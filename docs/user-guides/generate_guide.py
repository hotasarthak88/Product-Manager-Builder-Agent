"""Generate CLARA User Guide .docx"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import os

# Ensure we're in the right directory
os.chdir('/Users/sarthah/PMBuilderAgent')

doc = Document()

# --- Title Page ---
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('CLARA GRC Reviews AI Assistant', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('User Guide')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Author: Sarthak Hota (sarthah@)\n').font.size = Pt(12)
meta.add_run(f'Date: {date.today().strftime("%B %d, %Y")}\n').font.size = Pt(12)
meta.add_run('Version: v1.0\n').font.size = Pt(12)
meta.add_run('Audience: GRC Reviewers, GRC Approvers, GRC Controllership').font.size = Pt(12)

doc.add_page_break()

# --- Table of Contents ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. Accessing CLARA',
    '3. For GRC Reviewers',
    '   3.1 Open Your Review',
    '   3.2 Read CLARA Insights',
    '   3.3 Confirm Ownership',
    '   3.4 Step 1: Confirm Control Details',
    '   3.5 Step 2: Control Review Questions',
    '   3.6 Step 3: IPA Questions',
    '   3.7 Attach Evidence',
    '   3.8 Re-analyze and Validate',
    '   3.9 Submit',
    '4. For GRC Approvers',
    '   4.1 Open Review from Approval Queue',
    '   4.2 Check Review Status and SLA',
    '   4.3 Read CLARA Executive Summary',
    '   4.4 Review CLARA Assessments Per Question',
    '   4.5 Approve or Reject',
    '   4.6 Feedback Controls',
    '   4.7 Re-analyze Rate Limiting and Time-Outs',
    '5. Tips and Best Practices',
    '6. FAQ',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number' if not item.startswith('   ') else 'List Bullet')

doc.add_page_break()

# --- Section 1: Overview ---
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'CLARA (Compliance Learning and Review Assistant) is an AI-powered assistant integrated '
    'into the GRC 2.0 platform. It helps reviewers and approvers complete quarterly SOX control '
    'reviews faster and with fewer errors by providing:'
)
bullets = [
    'Pre-filled response recommendations based on historical patterns and control page data',
    'Real-time insights explaining why each answer is recommended',
    'Evidence validation — reads attached files and flags invalid or missing evidence',
    'Executive summaries for approvers with risk flags and key assertions',
    'Smart N/A detection for inapplicable questions based on control type',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'Important: CLARA assists with drafting but does not make decisions. All final responses '
    'are owned by the reviewer. AI content is fully suppressed once a review is submitted — '
    'auditors never see CLARA suggestions.'
).bold = False

doc.add_page_break()

# --- Section 2: Accessing CLARA ---
doc.add_heading('2. Accessing CLARA', level=1)
doc.add_paragraph(
    'CLARA is available on all control reviews in the GRC 2.0 platform (gamma environment). '
    'To access:'
)
doc.add_paragraph('1. Navigate to https://gamma.grc.a2z.com', style='List Number')
doc.add_paragraph('2. Open any review assigned to you from the review queue', style='List Number')
doc.add_paragraph('3. CLARA Insights panel appears automatically at the top of the review page', style='List Number')

doc.add_paragraph()
doc.add_paragraph(
    'No setup or configuration is required. CLARA activates automatically for all reviews in the pilot.'
)

doc.add_page_break()

# --- Section 3: For GRC Reviewers ---
doc.add_heading('3. For GRC Reviewers', level=1)
doc.add_paragraph(
    'This section walks you through the complete review workflow with CLARA assistance.'
)

# 3.1 Open Review
doc.add_heading('3.1 Open Your Review', level=2)
doc.add_paragraph(
    'Navigate to your review from the review queue. The page loads with two key sections visible:'
)
doc.add_paragraph('• CLARA Insights panel (top) — Executive Summary and Critical Context', style='List Bullet')
doc.add_paragraph('• Review content (below) — Ownership confirmation and questionnaire', style='List Bullet')
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/01-review-page-overview.png'):
    doc.add_picture('docs/user-guides/screenshots/01-review-page-overview.png', width=Inches(6))
    doc.add_paragraph('Figure 1: Review page overview showing CLARA Insights at top', style='Caption')

# 3.2 Read Insights
doc.add_heading('3.2 Read CLARA Insights', level=2)
doc.add_paragraph(
    'Before answering any questions, read the CLARA Insights panel at the top of the page. It contains:'
)
doc.add_paragraph('• Executive Summary — key findings, risks, and required actions', style='List Bullet')
doc.add_paragraph('• Critical Context — blocking issues that must be resolved before submission', style='List Bullet')
doc.add_paragraph('• "View Full Analysis" button — opens detailed per-question analysis in a modal', style='List Bullet')
doc.add_paragraph('• "Re-analyze" button — refreshes the analysis after you make changes', style='List Bullet')
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/03-insights-panel.png'):
    doc.add_picture('docs/user-guides/screenshots/03-insights-panel.png', width=Inches(6))
    doc.add_paragraph('Figure 2: CLARA Insights panel with Executive Summary and action buttons', style='Caption')

# 3.3 Confirm Ownership
doc.add_heading('3.3 Confirm Ownership', level=2)
doc.add_paragraph(
    'Before proceeding to the questionnaire, you must confirm that you are the owner of this review.'
)
doc.add_paragraph()
doc.add_paragraph('If you ARE the owner:', style='List Bullet')
doc.add_paragraph('  1. Select "Yes"', style='List Number')
doc.add_paragraph('  2. Click "Confirm and next" (orange button)', style='List Number')
doc.add_paragraph()
doc.add_paragraph('If you are NOT the owner:', style='List Bullet')
doc.add_paragraph('  1. Select "No"', style='List Number')
doc.add_paragraph('  2. Enter the correct owner\'s name in the transfer field', style='List Number')
doc.add_paragraph('  3. Click "Transfer"', style='List Number')
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/04-ownership-section.png'):
    doc.add_picture('docs/user-guides/screenshots/04-ownership-section.png', width=Inches(6))
    doc.add_paragraph('Figure 3: Ownership confirmation — select Yes/No and confirm or transfer', style='Caption')

# 3.4 Step 1
doc.add_heading('3.4 Step 1: Confirm Control Details', level=2)
doc.add_paragraph(
    'Confirm that the control details on the control page are accurate and up to date for the review period. '
    'CLARA provides a recommendation based on detected changes (or lack thereof).'
)
doc.add_paragraph()
doc.add_paragraph('• Review the CLARA recommendation panel below the question', style='List Bullet')
doc.add_paragraph('• Expand "Insights" to see why CLARA recommends this response', style='List Bullet')
doc.add_paragraph('• Accept the recommendation, edit it, or write your own response', style='List Bullet')
doc.add_paragraph('• Check the control page (right panel) to verify details yourself', style='List Bullet')
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/05-step1-confirm-details.png'):
    doc.add_picture('docs/user-guides/screenshots/05-step1-confirm-details.png', width=Inches(6))
    doc.add_paragraph('Figure 4: Step 1 with CLARA recommendation and control page details', style='Caption')

# 3.5 Step 2
doc.add_heading('3.5 Step 2: Control Review Questions', level=2)
doc.add_paragraph(
    'Answer each Yes/No question about control effectiveness, evidence, and compliance. '
    'CLARA provides per-question recommendations with supporting insights.'
)
doc.add_paragraph()
doc.add_paragraph('For each question:', style='List Bullet')
doc.add_paragraph('  1. Read the CLARA recommendation (Yes/No/N/A)', style='List Number')
doc.add_paragraph('  2. Expand "Insights" to understand the reasoning', style='List Number')
doc.add_paragraph('  3. Select your answer (you can disagree with CLARA)', style='List Number')
doc.add_paragraph('  4. Add comments if the question requires explanation', style='List Number')
doc.add_paragraph()
doc.add_paragraph(
    'Important: CLARA\'s recommendations are suggestions based on historical patterns. '
    'Always verify independently before accepting — especially for effectiveness questions.'
)
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/06-step2-questions.png'):
    doc.add_picture('docs/user-guides/screenshots/06-step2-questions.png', width=Inches(6))
    doc.add_paragraph('Figure 5: Step 2 questions with CLARA recommendations per question', style='Caption')

# 3.6 Step 3
doc.add_heading('3.6 Step 3: IPA Questions', level=2)
doc.add_paragraph(
    'IPA (Information Produced by Amazon) questions apply to manual and IT-dependent manual controls. '
    'For automated controls (IT General Controls), CLARA typically recommends N/A.'
)
doc.add_paragraph()
doc.add_paragraph('• If your control is automated: follow CLARA\'s N/A recommendation', style='List Bullet')
doc.add_paragraph('• If your control is manual with IPA: document source systems, queries, and owners', style='List Bullet')
doc.add_paragraph('• CLARA will flag if you select N/A when IPA documentation exists', style='List Bullet')
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/07-step3-ipa.png'):
    doc.add_picture('docs/user-guides/screenshots/07-step3-ipa.png', width=Inches(6))
    doc.add_paragraph('Figure 6: Step 3 IPA questions with CLARA recommendations', style='Caption')

# 3.7 Attach Evidence
doc.add_heading('3.7 Attach Evidence', level=2)
doc.add_paragraph(
    'Upload supporting evidence required by the test procedure. CLARA validates attached files — '
    'if a file doesn\'t match the expected evidence type, it will flag the issue.'
)
doc.add_paragraph()
doc.add_paragraph('1. Scroll to "Files & external link" section', style='List Number')
doc.add_paragraph('2. Click "Add attachment"', style='List Number')
doc.add_paragraph('3. Upload the required evidence (e.g., audit report, dashboard screenshot)', style='List Number')
doc.add_paragraph('4. Ensure the file name and description clearly identify the evidence', style='List Number')
doc.add_paragraph()
doc.add_paragraph(
    'Tip: After attaching evidence, click "Re-analyze" to have CLARA verify your attachment '
    'is valid and matches the test procedure requirements.'
)
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/08-files-attachments.png'):
    doc.add_picture('docs/user-guides/screenshots/08-files-attachments.png', width=Inches(6))
    doc.add_paragraph('Figure 7: Files & external link section for evidence upload', style='Caption')

# 3.8 Re-analyze
doc.add_heading('3.8 Re-analyze and Validate', level=2)
doc.add_paragraph(
    'After completing your responses and attaching evidence, click "Re-analyze" to have CLARA '
    'validate your submission. This will:'
)
doc.add_paragraph('• Check evidence validity (reads file content)', style='List Bullet')
doc.add_paragraph('• Identify contradictions between your answers', style='List Bullet')
doc.add_paragraph('• Update the Executive Summary with current state', style='List Bullet')
doc.add_paragraph('• Flag any remaining issues before submission', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph(
    'Review the updated Insights panel. If Critical Context shows issues, address them before submitting.'
)

# 3.9 Submit
doc.add_heading('3.9 Submit', level=2)
doc.add_paragraph(
    'Once all questions are answered, evidence is attached, and CLARA shows no critical flags:'
)
doc.add_paragraph('1. Verify all required fields are complete (Submit button enables when ready)', style='List Number')
doc.add_paragraph('2. Click "Submit" to send the review to your approver', style='List Number')
doc.add_paragraph()
doc.add_paragraph(
    'After submission, CLARA content is automatically hidden from the review. '
    'Your approver sees only your human-authored responses.'
)
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/09-submit-button.png'):
    doc.add_picture('docs/user-guides/screenshots/09-submit-button.png', width=Inches(6))
    doc.add_paragraph('Figure 8: Submit button (enabled when all required fields are complete)', style='Caption')

doc.add_page_break()

# --- Section 4: For GRC Approvers ---
doc.add_heading('4. For GRC Approvers', level=1)
doc.add_paragraph(
    'As an approver, CLARA provides you with an AI-generated summary and per-question assessments '
    'to help you make informed approval decisions quickly. This section walks you through the '
    'complete approval workflow.'
)

doc.add_heading('4.1 Open Review from Approval Queue', level=2)
doc.add_paragraph(
    'Open a submitted review from your approval queue. The page displays the reviewer\'s '
    'responses with CLARA analysis layered on top.'
)
doc.add_paragraph()
doc.add_paragraph('Key elements visible on the Approver view:', style='List Bullet')
doc.add_paragraph('• Review Status — shows reviewer name, submission date, and your due date', style='List Bullet')
doc.add_paragraph('• CLARA Insights panel — Executive Summary with critical findings', style='List Bullet')
doc.add_paragraph('• Reviewer responses — with CLARA assessment panels per question', style='List Bullet')
doc.add_paragraph('• Action buttons — "Approve" and "Reject with comment"', style='List Bullet')
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/10-approver-overview.png'):
    doc.add_picture('docs/user-guides/screenshots/10-approver-overview.png', width=Inches(6))
    doc.add_paragraph('Figure 9: Approver view — review page with approval controls', style='Caption')

doc.add_heading('4.2 Check Review Status and SLA', level=2)
doc.add_paragraph(
    'The Review Status section shows the approval chain with due dates for each approver level. '
    'Verify your approval deadline and note if the review was submitted on time.'
)
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/14-approver-review-status.png'):
    doc.add_picture('docs/user-guides/screenshots/14-approver-review-status.png', width=Inches(6))
    doc.add_paragraph('Figure 10: Review Status showing approval chain and due dates', style='Caption')

doc.add_heading('4.3 Read CLARA Executive Summary', level=2)
doc.add_paragraph(
    'The CLARA Insights panel at the top provides an AI-generated summary of the review quality. '
    'Read this FIRST — it tells you whether this review needs deep examination or is straightforward.'
)
doc.add_paragraph()
doc.add_paragraph('The summary includes:', style='List Bullet')
doc.add_paragraph('• Key findings — what CLARA identified as notable or problematic', style='List Bullet')
doc.add_paragraph('• Critical Context — blocking issues that may warrant rejection', style='List Bullet')
doc.add_paragraph('• Evidence assessment — whether attached files are valid and complete', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph(
    'Click "View Full Analysis" for a detailed per-question breakdown including confidence levels, '
    'historical patterns, and flags.'
)
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/11-approver-insights-summary.png'):
    doc.add_picture('docs/user-guides/screenshots/11-approver-insights-summary.png', width=Inches(6))
    doc.add_paragraph('Figure 11: CLARA Executive Summary showing critical findings for the approver', style='Caption')

doc.add_heading('4.4 Review CLARA Assessments Per Question', level=2)
doc.add_paragraph(
    'Each question the reviewer answered has a "CLARA assessment" expandable panel. '
    'Click to expand and see CLARA\'s evaluation of the reviewer\'s response.'
)
doc.add_paragraph()
doc.add_paragraph('Each assessment shows:', style='List Bullet')
doc.add_paragraph('• Whether the response aligns with historical patterns', style='List Bullet')
doc.add_paragraph('• Confidence level (High / Medium / Low)', style='List Bullet')
doc.add_paragraph('• Flags if the response contradicts evidence or control page data', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph(
    'Tip: Focus on questions where CLARA shows Low confidence or flags — these are the '
    'ones most likely to have issues.'
)
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/12-approver-clara-assessment.png'):
    doc.add_picture('docs/user-guides/screenshots/12-approver-clara-assessment.png', width=Inches(6))
    doc.add_paragraph('Figure 12: CLARA assessment panels on reviewer responses (expandable)', style='Caption')

doc.add_heading('4.5 Approve or Reject', level=2)
doc.add_paragraph(
    'After reviewing the summary and responses, make your decision:'
)
doc.add_paragraph()
doc.add_paragraph('To Approve:', style='List Bullet')
doc.add_paragraph('  • Click "Approve" — the review moves to the next approval level or completes', style='List Bullet')
doc.add_paragraph('  • Only approve if you\'re satisfied that responses are accurate and evidence is valid', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('To Reject:', style='List Bullet')
doc.add_paragraph('  • Click "Reject with comment"', style='List Bullet')
doc.add_paragraph('  • Write specific feedback explaining what needs to be fixed', style='List Bullet')
doc.add_paragraph('  • Reference CLARA\'s findings to provide actionable rejection reasons', style='List Bullet')
doc.add_paragraph('  • The review returns to the reviewer as a new draft with your comments', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph(
    'Best practice: When rejecting, be specific. Instead of "fix the evidence," write '
    '"Replace \'evidence ppt\' (test file) with the actual quarterly FARM report showing control operation for this period."'
)
doc.add_paragraph()
if os.path.exists('docs/user-guides/screenshots/13-approver-approve-reject.png'):
    doc.add_picture('docs/user-guides/screenshots/13-approver-approve-reject.png', width=Inches(6))
    doc.add_paragraph('Figure 13: Approve and Reject with comment buttons', style='Caption')

doc.add_page_break()

# --- Section 4.6: Feedback Controls ---
doc.add_heading('4.6 Feedback Controls', level=2)
doc.add_paragraph(
    'Each CLARA recommendation includes three feedback controls at the bottom of the insight card:'
)
doc.add_paragraph()
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
# Header
table.rows[0].cells[0].text = 'Control'
table.rows[0].cells[1].text = 'Icon'
table.rows[0].cells[2].text = 'When to Use'
# Row 1
table.rows[1].cells[0].text = 'Thumbs Up (Helpful)'
table.rows[1].cells[1].text = '👍'
table.rows[1].cells[2].text = 'Click when the recommendation or insight was accurate and useful. This helps CLARA learn which responses work well for your control type.'
# Row 2
table.rows[2].cells[0].text = 'Thumbs Down (Not helpful)'
table.rows[2].cells[1].text = '👎'
table.rows[2].cells[2].text = 'Click when the recommendation was inaccurate, irrelevant, or unhelpful. This signals CLARA to improve its suggestions for similar controls.'
# Row 3
table.rows[3].cells[0].text = 'Copy'
table.rows[3].cells[1].text = '📋'
table.rows[3].cells[2].text = 'Click to copy the CLARA recommendation text to your clipboard. Useful when you want to paste the suggestion into the response field and then edit it.'

doc.add_paragraph()
doc.add_paragraph(
    'Your feedback is anonymous and used solely to improve CLARA\'s recommendations. '
    'It does not affect your review submission or approval status.'
)

doc.add_page_break()

# --- Section 4.7: Re-analyze Rate Limiting ---
doc.add_heading('4.7 Re-analyze: Rate Limiting and Time-Outs', level=2)
doc.add_paragraph(
    'CLARA enforces rate limiting on the Re-analyze function to manage system resources. '
    'You may encounter a time-out under these conditions:'
)
doc.add_paragraph()
doc.add_paragraph(
    '1. Multiple browser sessions are open for the same review and Re-analyze is clicked across those sessions simultaneously.',
    style='List Number'
)
doc.add_paragraph(
    '2. Re-analyze is clicked multiple times within a 5-minute window on the same review.',
    style='List Number'
)
doc.add_paragraph()
doc.add_paragraph('What happens when rate-limited:')
doc.add_paragraph('• CLARA displays a time-out message indicating the analysis cannot run right now', style='List Bullet')
doc.add_paragraph('• The time-out typically lasts 5–10 minutes', style='List Bullet')
doc.add_paragraph('• Your existing analysis (from the last successful run) remains visible and valid', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('What to do:')
doc.add_paragraph('• Wait 5–10 minutes before clicking Re-analyze again', style='List Bullet')
doc.add_paragraph('• Close duplicate browser tabs/sessions for the same review', style='List Bullet')
doc.add_paragraph('• Continue working on your responses — the rate limit only affects Re-analyze, not your ability to answer questions or submit', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph(
    'This is expected behavior, not an error. If the time-out persists beyond 15 minutes, '
    'contact gfrc-sox@amazon.com for assistance.'
)

doc.add_page_break()

# --- Section 5: Tips ---
doc.add_heading('5. Tips and Best Practices', level=1)
tips = [
    ('Always click Re-analyze after changes', 'CLARA analysis is cached. After attaching evidence or changing answers, click Re-analyze to get updated validation.'),
    ('Don\'t blindly accept recommendations', 'CLARA bases suggestions on historical patterns. If your situation is different this quarter, override the recommendation.'),
    ('Check the Full Analysis for flags', 'The inline insights are brief. Click "View Full Analysis" for detailed per-question analysis including confidence levels and flags.'),
    ('Attach evidence BEFORE submitting', 'CLARA validates evidence content. Attaching the right evidence early prevents approver rejection.'),
    ('Review Critical Context first', 'If there are blocking issues, address them before spending time on questions.'),
]
for title_text, desc in tips:
    p = doc.add_paragraph()
    p.add_run(f'{title_text}: ').bold = True
    p.add_run(desc)

doc.add_page_break()

# --- Section 6: FAQ ---
doc.add_heading('6. FAQ', level=1)
faqs = [
    ('Can auditors see CLARA suggestions?', 'No. Once a review is submitted, all AI content is automatically suppressed. Auditors only see human-authored responses.'),
    ('What if CLARA is wrong?', 'CLARA provides recommendations, not mandates. You can always override with your own response. Use the "Discard" option to write from scratch.'),
    ('Why does CLARA say "Generated by Clara — AI-generated content may be incorrect"?', 'This is a standard disclaimer. CLARA assists with drafting but may not have complete context. Always verify independently.'),
    ('What happens if I don\'t click Re-analyze?', 'The CLARA analysis remains cached from the last run. New evidence or answer changes won\'t be reflected in the Insights panel until you re-analyze.'),
    ('Can I use CLARA for all control types?', 'CLARA is available for all control types. Per-question recommendations may vary in availability depending on the control type and historical data.'),
    ('Who sees my review before I submit?', 'Only you. CLARA assists you privately during the draft phase. After submission, only your approver sees the final responses.'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    p.add_run(f'Q: {q}').bold = True
    doc.add_paragraph(f'A: {a}')
    doc.add_paragraph()

# --- Save ---
output_path = 'docs/user-guides/CLARA-GRC-Reviews-AI-Assistant-User-Guide-v1.0.docx'
doc.save(output_path)
print(f'Guide saved to: {output_path}')
